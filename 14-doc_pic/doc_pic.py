from docx import Document
from docx.shared import Inches


doc = Document()
# 添加图片
doc.add_picture("1.jpg", width=Inches(1.25))
doc.save("new_pic.docx")


doc = Document()
# 创建table
table = doc.add_table(rows=3, cols=4)
# 设置table样式
table.style = "Table Grid"

# 第一种方法 先获取行 再获取该行中对应的单元格
row = table.rows[0]
row.cells[0].text = "第一行第一列"

# 第二种方法 直接指行号和列号
cell = table.cell(0, 1)
cell.text = "第一行第一列"

# 获取表格中的单元格对象
cell = table.cell(1, 0)
# 获取单元格中的段落对象
p = cell.paragraphs[0]
# 获取追加对象
run = p.add_run()
run.add_picture("1.jpg", width=Inches(1.25))

doc.save("new_rows.docx")
