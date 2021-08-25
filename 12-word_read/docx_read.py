from docx import Document


# 创建文档对象
document = Document()
# 保存文档对象 扩展名只可以使用*.docx
document.save("new.docx")

# 获取文档
doc = Document("exist.docx")
# 遍历Word文档中的段落
for p in doc.paragraphs:
    # 输出word文档中的段落内容
    print(p.text)

print("===============")

# 拿到表格
tables = doc.tables
table = tables[0]
values = []
# 遍历表格的每一行
for row in table.rows:
    # 遍历每一行中的单元格
    for cell in row.cells:
        # 将单元格中的数据添加到list中
        values.append(cell.text)
    value = ' '.join(values)
    print(value)
    values = []

