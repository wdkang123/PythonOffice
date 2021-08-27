import pandas as pd
import matplotlib.pyplot as plt
from docx import Document


# 自动生成数据分析报告
students = pd.read_excel("student_score.xlsx")
# 就地排序
students.sort_values(by='Score', inplace=True, ascending=False)
# 绘制柱状图
plt.bar(students["Name"], students["Score"], color="orange")

plt.title("Student Score", fontsize=16)
plt.xlabel("Name")
plt.ylabel("Score")

# 重铺x轴标签
plt.xticks(students.Name, rotation='90')
# 紧凑型布局
plt.tight_layout()
# 保存图片
plt.savefig("student_score.png")

# 数据虽然被排序 但数据下标并没有改变 直接通过下标获取的依旧是原本的数据
print("原始下标: ", students["Name"][0])

# 方法一 通过iloc方法获取绝对的位置
print("绝对位置: ", students.iloc[0, :]["Name"])
# 方法二 对下标重新排序再取其中第一位
students.reset_index(drop=True, inplace=True)
print("重新排序后的下标 ", students["Name"][0])


# 生成word文档
doc = Document()
doc.add_heading("数据分析报告", level=0)
# 绝对定位 获取分数排在第一位的学生信息
first_student = students.iloc[0, :]["Name"]
first_score = students.iloc[0, :]["Score"]

p = doc.add_paragraph("分数排在第一位的学生是: ")
# 设置为粗体
p.add_run(str(first_student)).bold = True
p.add_run(', 分数为 ')
p.add_run(str(first_score)).bold = True

p1 = doc.add_paragraph(f"总共有 {len(students['Name'])} 名学生参加了考试 学生考试总体情况为 ")

# 添加表格
table = doc.add_table(rows=len(students["Name"]) + 1, cols=2)

# 设置表格样式
table.style = "LightShading-Accent1"
table.cell(0, 0).text = "学生姓名"
table.cell(0, 1).text = "学生分数"

# 添加数据到表中
for i, (index, row) in enumerate(students.iterrows()):
  table.cell(i + 1, 0).text = str(row["Name"])
  table.cell(i + 1, 1).text = str(row["Score"])

# 添加图片
doc.add_picture("student_score.png")
doc.save("student_score_analyze.docx")
print("Done!")
