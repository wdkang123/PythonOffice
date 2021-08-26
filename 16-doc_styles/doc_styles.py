from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import *
from docx.shared import Pt


# 修改word文档样式
doc = Document()

p1 = doc.add_paragraph("水平居中对齐")
# 设置段落水平居中对齐
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph("左对齐")
# 设置段落左对齐
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

p3 = doc.add_paragraph("右对齐")
# 设置段落右对齐
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p4 = doc.add_paragraph()
run = p4.add_run("内联对象")
font = run.font
# 设置字体大小
font.size = Pt(35)
# 设置字体为斜体
font.italic = True


# 获取python-docx支持的所有样式
styles = doc.styles
# 选取style并设置style中的段落格式
style = styles["Heading 1"]
p_format = style.paragraph_format
# 设置左缩进
p_format.left_indent = Pt(25)
# 使用样式
p = doc.add_paragraph("使用style设置段落样式", style=style)

doc.save("doc_style.docx")


# 输出所有表格的样式
doc = Document()
styles = doc.styles
for style in styles:
  # 过滤表格样式
  if style.type == WD_STYLE_TYPE.TABLE:
    # 输出当前样式的样式名
    doc.add_paragraph(f"表格样式名称: {style.name}")
    # 创建表格并指定为当前样式
    table = doc.add_table(3, 3, style=style)
    # 将内容添加到第一行
    cells = table.rows[0].cells
    cells[0].text = "第一列内容"
    cells[1].text = "第二列内容"
    cells[2].text = "第三列内容"
    doc.add_paragraph("\n")

doc.save("doc_all_style_list.docx")
