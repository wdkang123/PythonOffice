from docx import Document


doc = Document()

# 有序列表
style = "List Number"
doc.add_paragraph("有序列表1", style=style)
doc.add_paragraph("有序列表2", style=style)
doc.add_paragraph("有序列表3", style=style)

# 无序列表
style = "List Bullet"
doc.add_paragraph("无序列表1", style=style)
doc.add_paragraph("无序列表2", style=style)
doc.add_paragraph("无序列表3", style=style)

doc.save("doc_list.docx")
