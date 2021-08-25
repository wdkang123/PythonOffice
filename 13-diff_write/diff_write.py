from docx import Document
from deepdiff import DeepDiff


def get_doc_values(path):
    # 获取表格中的内容 返回二维数组
    doc = Document(path)
    tables = doc.tables
    table = tables[0]
    all_values = []
    for row in table.rows:
        values = []
        for cell in row.cells:
            # 将表格中单元格的值添加到values列表中
            values.append(cell.text)
        all_values.append(values)
    return all_values


table1 = get_doc_values("table.docx")
table2 = get_doc_values("table_modify.docx")

# 比较列表差异
ddiff = DeepDiff(table1, table2)
print(ddiff)

'''
{'values_changed': {'root[2][1]': {'new_value': '1973', 'old_value': '1972'}, 'root[3][0]': 
{'new_value': 'JVAA', 'old_value': 'JAVA'}, 'root[3][1]': {'new_value': '1996', 'old_value': '1995'}}}
'''

# 写入文档
doc = Document()
# 添加标题
doc.add_heading("一级标题", level=1)
# 添加段落
p2 = doc.add_paragraph("第二个段落")
# 将新段落添加到已经有的段落之前
p1 = p2.insert_paragraph_before("第一个段落")

p3 = doc.add_paragraph("新段落")
# 追加内容
p3.add_run("加粗").bold = True
p3.add_run("以及")
p3.add_run("斜体").italic = True

doc.save("new_doc.docx")
